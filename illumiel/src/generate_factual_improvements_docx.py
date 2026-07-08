"""
일루미엘 20대 팩트 기반 개선사항 Word 생성 스크립트
작성자: 데이터 분석가
목적: 100% 사실 데이터에만 근거한 20개의 핵심 비즈니스 액션 플랜을 도출하고 공식 .docx 파일로 저장
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# Set paths
base_dir = r'c:\Users\다의\Desktop\icb10proj2\illumiel'
report_dir = os.path.join(base_dir, 'report')
docx_path = os.path.join(report_dir, '일루미엘_6월_팩트기반_20대개선사항.docx')

document = Document()

# Title
title = document.add_heading('일루미엘 6월 팩트 기반 20대 핵심 개선사항', 0)
title.style.font.name = 'Malgun Gothic'
title.style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Intro
p_intro = document.add_paragraph('본 문서는 상상력과 허위 지표를 완전히 배제하고, 오직 6월 스토어의 실제 팩트 데이터(매출액, 고유 방문자수, 전환율, 객단가 등)에 기반하여 작성된 20가지 비즈니스 액션 플랜입니다. 모든 핵심 문장의 끝에는 도출 근거가 되는 수치가 명시되어 있습니다.')

document.add_heading('■ 트래픽 방어 및 UI/UX 개선', level=1)

improvements = [
    ("대규모 기획전 진행 시 상세페이지 이미지 용량을 최적화(WebP 등)하여 로딩 지연으로 인한 고객 이탈을 철저히 방어해야 합니다.", "(고유 방문 고객수 20,336명 폭증 근거)"),
    ("폭발적으로 늘어난 트래픽이 허수로 빠져나가지 않도록 스토어 내 메인 배너 및 카테고리의 직관성을 대폭 개선해야 합니다.", "(전월 대비 방문 고객수 +193.15% 증가폭 근거)"),
    ("일시적인 방문자 폭증 상황에서도 결제 심리가 무너지지 않도록, 결제 페이지 진입 시 로딩 없는 쾌적한 환경을 유지해야 합니다.", "(스토어 평균 구매 전환율 2.31% 방어 근거)"),
    ("핵심 매출을 견인하는 주력 상품의 트래픽을 타겟으로, 상세페이지 최상단 3초 체류 구간에 비포/애프터 리뷰를 고정 배치하여 설득력을 높여야 합니다.", "(감태/레스큐 상품군 방문수 37,509회 발생 근거)"),
    ("방문수가 압도적으로 높은 주력 상품의 체류 시간을 늘리기 위해, 텍스트 설명을 최소화하고 가독성 높은 영상 위주의 콘텐츠로 리뉴얼해야 합니다.", "(감태/레스큐 방문수 37,509회 근거)")
]

for text, basis in improvements:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text + " ")
    run_basis = p.add_run(basis)
    run_basis.bold = True
    run_basis.font.color.rgb = RGBColor(0, 112, 192)

document.add_heading('■ 수익성(객단가) 강화 전략', level=1)

improvements2 = [
    ("객단가가 낮은 특정 상품의 상세페이지 하단에 고단가 베스트셀러를 함께 노출시켜 자연스러운 교차 구매(Cross-selling)를 유도해야 합니다.", "(신규 고객 79.7% 유입 및 객단가 19,668원 근거)"),
    ("저단가 체리피킹 트래픽을 수익으로 치환하기 위해, 단품 구매 시 마진이 높은 크림을 엮은 이종 결합 번들(Bundle) 상품을 신설해야 합니다.", "(클리어런스 결제건수 300건 대비 판매금액 2,585,420원 근거)"),
    ("수익의 절대 다수를 차지하는 핵심 캐시카우 상품 라인업이 외부 환경 변화에 흔들리지 않도록 자사 스토어 찜/알림받기 유도를 강화해야 합니다.", "(감태/레스큐 상품군 판매금액 5,265,410원 근거)"),
    ("적은 유입으로도 높은 수익을 올리는 외부 인플루언서 공구 마케팅을 정례화하여, 효율성 높은 타겟 집중형 매출 구조를 추가 확보해야 합니다.", "(나나블리 특가 방문수 326회 대비 판매금액 1,220,100원 달성 근거)"),
    ("고관여 고객층의 트래픽을 직접 스토어 자산으로 만들기 위해, 인플루언서 랜딩 페이지에 '알림받기 필수' 조건부 혜택 팝업을 삽입해야 합니다.", "(나나블리 특가 결제건수 25건 근거)")
]

for text, basis in improvements2:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text + " ")
    run_basis = p.add_run(basis)
    run_basis.bold = True
    run_basis.font.color.rgb = RGBColor(0, 112, 192)

document.add_heading('■ 리텐션(재구매) 극대화 및 CRM', level=1)

improvements3 = [
    ("신규로 대거 유입된 구매자들을 대상으로 구매일 기준 20일 뒤 '시크릿 재구매 20% 쿠폰' 알림톡을 발송하여 리텐션을 극대화해야 합니다.", "(신규 고객 결제 객단가 19,668원 근거)"),
    ("전체 매출의 큰 비중을 차지하는 신규 고객의 2차 구매를 유도하기 위해, 출고 시 첫 구매 감사 인사와 샘플을 동봉하는 웰컴 패키징을 도입해야 합니다.", "(신규 고객 6월 매출의 79.7% 차지 근거)"),
    ("브랜드의 수익성을 든든하게 받쳐주는 기존 고객을 위해 VIP 전용 '신제품 선행 샘플 증정' 등 프리미엄 멤버십 제도를 신설해야 합니다.", "(재구매 고객 구매 전환율 23.92% 및 객단가 42,910원 근거)"),
    ("압도적인 재구매 전환율을 지닌 기존 고객들의 충성도를 높이기 위해, 고객 피드백을 적극 수용한 '감태 리뉴얼 한정 에디션' 출시를 기획해야 합니다.", "(재구매 고객 구매 전환율 23.92% 근거)"),
    ("재구매 확률과 객단가가 월등히 높은 충성 고객층을 핵심 타겟으로, 고단가 기초세트의 '매월 정기 결제(구독)' 모델 런칭을 서둘러야 합니다.", "(재구매 고객 객단가 42,910원 근거)")
]

for text, basis in improvements3:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text + " ")
    run_basis = p.add_run(basis)
    run_basis.bold = True
    run_basis.font.color.rgb = RGBColor(0, 112, 192)

document.add_heading('■ 운영 효율화 및 기타 비즈니스 액션', level=1)

improvements4 = [
    ("매출 외형 성장이 두드러진 만큼, 배송 지연이나 단순 변심으로 인한 취소 및 환불을 막기 위해 '출고 전 감성 알림톡' 발송 프로세스를 즉시 구축해야 합니다.", "(전월 대비 총 매출액 +58.45% 폭풍 성장 근거)"),
    ("결제 건수 증가폭이 큰 만큼 물류 포장 및 배송 CS 인력을 단기 확충하여 고객 불만을 사전 차단하고 긍정적 포토 리뷰를 양산해야 합니다.", "(전월 대비 결제건수 +166.83% 증가 근거)"),
    ("재고 소진 목적의 기획전에서 확보된 높은 결제 행동 데이터를 전략적으로 활용하여, 클리어런스 페이지 이탈자 대상 장바구니 리타겟팅을 세팅해야 합니다.", "(클리어런스 상품군 결제건수 300건 및 구매전환율 17.21% 근거)"),
    ("압도적으로 높은 전환율을 보이는 재고 상품들의 특성을 분석하여, 향후 '48시간 타임세일 한정 특가' 형태로 트래픽을 순간적으로 모으는 기획을 상설화해야 합니다.", "(클리어런스 상품군 전환율 17.21% 근거)"),
    ("스토어 유입 트래픽 대비 실제 결제 건수가 안정적인 우상향 곡선을 그릴 수 있도록, 공백기 없는 촘촘한 평시 노출 기획전을 월간 캘린더에 편성해야 합니다.", "(총 6월 결제건수 539건 달성 근거)")
]

for text, basis in improvements4:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text + " ")
    run_basis = p.add_run(basis)
    run_basis.bold = True
    run_basis.font.color.rgb = RGBColor(0, 112, 192)

document.save(docx_path)
print(f"팩트 기반 20대 개선사항 Word 파일 생성 완료: {docx_path}")
